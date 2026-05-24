from pathlib import Path

from pypdf import PdfReader


SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".xlsx"}


def extract_text_from_file(path: Path, filename: str) -> str:
    """Extract plain text from supported office documents."""
    ext = Path(filename).suffix.lower()
    if ext == ".pdf":
        return _pdf_text(path)
    if ext == ".docx":
        return _docx_text(path)
    if ext == ".txt":
        return path.read_text(encoding="utf-8", errors="ignore")
    if ext == ".xlsx":
        return _xlsx_text(path)
    raise ValueError(f"Unsupported file type: {ext}")


def extract_pages_from_file(path: Path, filename: str) -> list[tuple[int, str]]:
    """Return (page_number, text) for chunking. Non-PDF types use a single virtual page."""
    text = extract_text_from_file(path, filename).strip()
    if not text:
        return []
    return [(1, text)]


def _pdf_text(path: Path) -> str:
    reader = PdfReader(str(path))
    parts = []
    for page in reader.pages:
        parts.append(page.extract_text() or "")
    return "\n\n".join(parts)


def _docx_text(path: Path) -> str:
    from docx import Document

    doc = Document(str(path))
    parts: list[str] = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n\n".join(parts)


def _xlsx_text(path: Path) -> str:
    from openpyxl import load_workbook

    wb = load_workbook(str(path), read_only=True, data_only=True)
    parts: list[str] = []
    for sheet in wb.worksheets:
        parts.append(f"## Sheet: {sheet.title}")
        for row in sheet.iter_rows(values_only=True):
            cells = [str(c).strip() for c in row if c is not None and str(c).strip()]
            if cells:
                parts.append(" | ".join(cells))
    wb.close()
    return "\n\n".join(parts)
